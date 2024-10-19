from docx import Document
from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, RGBColor, Cm
import docx


class Word():
    def __init__(self, title_document, header_document, number_header_document, exam_number, questions, number_questions, solution, options_supported, destination, file_name) -> None:
        self.title_document = title_document
        self.header_document = header_document
        self.number_header_document = number_header_document
        self.exam_number = exam_number
        self.questions = questions
        self.number_questions = number_questions
        self.solution = solution
        self.options_supported = options_supported
        self.destination = destination
        self.file_name = file_name
        self.doc = None
        self._set_word_document()
        self._write_word_document()
        self._save_word_document()

        
    def _set_word_document(self) -> None:
        self.doc = Document()

        # Imposta il font del documento
        styles = self.doc.styles['Normal']
        font = styles.font
        font.name = 'Liberation Sans'
        
        # Imposta la lingua del documento
        styles_element = self.doc.styles.element
        rpr_default = styles_element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0]
        lang_default = rpr_default.xpath('w:lang')[0]
        lang_default.set(docx.oxml.shared.qn('w:val'),'it-IT')
        
        # Imposta il titolo del documento
        par = self.doc.add_paragraph(self.title_document + f" - #{self.exam_number}") if self.number_header_document else self.doc.add_paragraph(self.title_document)
        for run in par.runs:
            run.bold = True
            run.font.size = Pt(15)
        par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Imposta i margini del documento
        section = self.doc.sections[0]
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
        
        # Imposta l'intestazione del documento
        upper = section.header
        paragraph = upper.paragraphs[0]
        paragraph.text = self.header_document        

        ## Aggiunge numero di pagina
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        self.doc.sections[0].footer.paragraphs[0].add_run()._r.append(fldChar1)
        self.doc.sections[0].footer.paragraphs[0].add_run()._r.append(instrText)
        self.doc.sections[0].footer.paragraphs[0].add_run()._r.append(fldChar2)
        
        # Aggiunge nuova sezione con due colonne
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        section = self.doc.sections[1]
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')


    def _write_word_document(self) -> None:
        for index, question in enumerate(self.questions):
            header_question = (f"{index + 1}) " if self.number_questions else "") + question['question']
            
            # Aggiunge il titolo della domanda
            h = self.doc.add_heading(header_question, 3)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = True
                run.font.name = 'Liberation Sans'
                run.font.size = Pt(11)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            # Aggiunge le opzioni alla domanda
            for i in range(0, self.options_supported):
                p = self.doc.add_paragraph(style='List Bullet')
                r = p.add_run(question['options'][i]['text'])
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                if self.solution:
                    r.bold = question['options'][i]['correct']
                    r.underline = question['options'][i]['correct']

                
    def _save_word_document(self) -> None: 
        suffix = "_solutions" if self.solution else ""
        self.doc.save(f"{self.destination}/{self.file_name}_{str(self.exam_number)}{suffix}.docx")    


