from docx import Document
from docx.document import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, ns
from docx.shared import Pt, RGBColor, Cm
import docx
from docx.shared import Inches
from configs import Configuration


class Word():
    def __init__(self, questions: list, document_number: int, is_solution: bool) -> None:
        self.questions = questions
        self.document_number = document_number
        self.is_solution = is_solution
        self.doc = Document()
        self.set_document()

    def set_document(self) -> None:
        self.set_font()
        self.set_language()
        self.set_margin()

        if Configuration.get_is_header_included():
            self.set_header()
        
        self.set_title()

        if Configuration.get_is_subtitle_included():
            self.set_subtitle()

        if Configuration.get_are_pages_numbered():
            self.set_number_page()

        self.set_columns()
        self.write_questions()

    def set_font(self) -> None:
        self.doc.styles['Normal'].font.name = Configuration.get_font()

    def set_language(self) -> None:
        self.doc.styles.element.xpath('./w:docDefaults/w:rPrDefault/w:rPr')[0].xpath('w:lang')[0].set(docx.oxml.shared.qn('w:val'), Configuration.get_language())

    def set_title(self) -> None:
        paragraph = self.doc.add_paragraph(Configuration.get_document_title() + f" - #{self.document_number}") if Configuration.get_are_documents_numbered() else self.doc.add_paragraph(Configuration.get_document_title())
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(Configuration.get_title_size())
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_subtitle(self) -> None:
        paragraph = self.doc.add_paragraph(Configuration.get_document_subtitle())
        for run in paragraph.runs:
            run.italic = True
            run.font.size = Pt(Configuration.get_subtitle_size())
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def set_margin(self) -> None:
        self.doc.sections[0].left_margin = Cm(Configuration.get_left_margin())
        self.doc.sections[0].right_margin = Cm(Configuration.get_right_margin())

    def set_number_page(self) -> None:
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'end')
        self.doc.sections[0].footer.paragraphs[0].add_run()._r.append(fldChar1)
        self.doc.sections[0].footer.paragraphs[0].add_run()._r.append(instrText)
        self.doc.sections[0].footer.paragraphs[0].add_run()._r.append(fldChar2)

    def set_header(self) -> None:
        self.doc.sections[0].header.paragraphs[0].text = Configuration.get_document_header()

    def set_columns(self) -> None:
        self.doc.add_section(WD_SECTION.CONTINUOUS)
        self.doc.sections[1]._sectPr.xpath('./w:cols')[0].set(qn('w:num'), str(Configuration.get_columns_number()))

    def add_image_shape(self, image: object) -> None:
        # Accedi all'elemento XML dell'immagine
        inline_shape = image._inline
        # Ottieni l'elemento grafico dell'immagine
        graphic = inline_shape.xpath('./a:graphic')[0]
        graphic_data = graphic.xpath('./a:graphicData')[0]
        pic = graphic_data.xpath('./pic:pic')[0]

        # Accedi all'elemento "spPr" per modificare le proprietà di stile
        spPr = pic.xpath('./pic:spPr')[0]
        ln = OxmlElement('a:ln')
        # Spessore del bordo (12700 EMU equivale a circa 1pt)
        ln.set('w', '12700')
        ln.set('cap', 'flat')

        # Colore del bordo (es. grigio chiaro)
        solid_fill = OxmlElement('a:solidFill')
        srgb_clr = OxmlElement('a:srgbClr')
        srgb_clr.set('val', '505050')  # Colore esadecimale grigio chiaro
        solid_fill.append(srgb_clr)
        ln.append(solid_fill)

        # Aggiungi l'elemento "ln" al "spPr"
        spPr.append(ln)

    def add_question_header(self, question_header: str, question_image: str) -> None:
        text_header = f"{question_header}\n" if question_image else f"{question_header}"
        header = self.doc.add_heading(text_header, 3)
        if question_image:
            run_i = header.add_run()
            image = run_i.add_picture(question_image, width=Inches(Configuration.get_images_size()))
            self.add_image_shape(image)

        if not question_image:
            header.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        for run in header.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
            run.font.name = Configuration.get_font()
            run.font.size = Pt(Configuration.get_questions_size())

    def add_question_option(self, option: dict) -> None:
        paragraph = self.doc.add_paragraph(option['text'], style='List Bullet')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        
        for run in paragraph.runs:
            # run.font.size = 18  #TODO
            if self.is_solution:
                run.bold = option["correct"]
                run.underline = option["correct"]
                if option["correct"]:
                    run.font.color.rgb = RGBColor(15, 150, 0)

    def write_questions(self) -> None:
        for index, question in enumerate(self.questions):
            question_header = (f"{index + 1}) " if Configuration.get_are_questions_numbered() else "") + question['question']
            question_image = question['image']

            self.add_question_header(question_header, question_image)
            for option in question['options']:
                self.add_question_option(option)

    def save(self) -> str:
        suffix = "_solutions" if self.is_solution else ""
        document_path = f"{Configuration.get_documents_directory()}/{Configuration.get_document_filename()}_{str(self.document_number)}{suffix}.docx"
        self.doc.save(document_path)
        return document_path
